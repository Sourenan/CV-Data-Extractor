"""
Word document parser implementation using python-docx.

python-docx gives direct access to paragraph-level text, which is
cleaner than trying to interpret raw XML. Table cells are also included
because resumes often use tables for skills or contact info.
"""

import os

from docx import Document
from docx.oxml.ns import qn

from app.exceptions import FileParsingError
from app.parsers.base import FileParser
from app.utils.logger import get_logger
from app.utils.text_utils import normalize_whitespace

logger = get_logger(__name__)


def _extract_table_text(doc: Document) -> list[str]:
    """Return all non-empty text cells from every table in *doc*."""
    cell_texts: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    cell_texts.append(cell_text)
    return cell_texts


class WordParser(FileParser):
    """Extracts plain text from a ``.docx`` Word document."""

    def parse(self, file_path: str) -> str:
        """Read *file_path* and return its full text content.

        Paragraphs are joined with newlines; table cells are appended
        after paragraph text so that skills listed in tables are still
        visible to the extractors.

        Parameters
        ----------
        file_path:
            Path to a ``.docx`` file.

        Returns
        -------
        str
            Normalised text extracted from all paragraphs and tables.
            Empty string if the document contains no text.

        Raises
        ------
        FileParsingError
            If the file does not exist, is not a valid ``.docx``, or
            python-docx raises an unrecoverable error.
        """
        if not os.path.isfile(file_path):
            raise FileParsingError(file_path, "file not found")

        logger.info("Parsing Word document: %s", file_path)

        try:
            doc = Document(file_path)
        except Exception as exc:
            raise FileParsingError(
                file_path, f"could not open Word document: {exc}"
            ) from exc

        try:
            paragraph_lines = [p.text for p in doc.paragraphs]
            table_lines = _extract_table_text(doc)
        except Exception as exc:
            raise FileParsingError(
                file_path, f"error reading document content: {exc}"
            ) from exc

        raw_text = "\n".join(paragraph_lines + table_lines)
        text = normalize_whitespace(raw_text)

        if not text:
            logger.warning(
                "Word document produced no extractable text: %s", file_path
            )

        logger.debug(
            "Extracted %d characters from Word document: %s", len(text), file_path
        )
        return text
